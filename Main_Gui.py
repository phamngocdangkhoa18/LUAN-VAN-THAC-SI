# %%
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog
import cv2
import pyautogui
import pyperclip

# Xử lý đa luồng
class WorkerThread(QtCore.QThread):
    result_ready = QtCore.pyqtSignal(object, str, object)
    error_occurred = QtCore.pyqtSignal(str)

    def __init__(self, image_path, selected_struct="", real_width_cm=0, real_height_cm=0, mode="overview"):
        super().__init__()
        self.image_path = image_path
        self.selected_struct = selected_struct
        self.real_width_cm = real_width_cm
        self.real_height_cm = real_height_cm
        self.mode = mode  # "overview" hoặc "zoom"

    def run(self):
        try:
            if self.mode == "overview":
                from crack_pipeline import run_pipeline
                import cv2
                image = cv2.imread(self.image_path)
                result_image, result_text, chart_np = run_pipeline(image)
            else:
                from crack_pipeline import run_pipeline2
                result_image, result_text, chart_np = run_pipeline2(
                    self.image_path,
                    self.selected_struct,
                    self.real_width_cm,
                    self.real_height_cm
                )

            self.result_ready.emit(result_image, result_text, chart_np)
        except Exception as e:
            self.error_occurred.emit(str(e))


# zoom ảnh nhấp vào
class ZoomableImageViewer(QtWidgets.QGraphicsView):
    def __init__(self, pixmap, parent=None):
        super().__init__(parent)
        self.scene = QtWidgets.QGraphicsScene(self)
        self.pixmap_item = QtWidgets.QGraphicsPixmapItem(pixmap)
        self.scene.addItem(self.pixmap_item)
        self.setScene(self.scene)
        self.setRenderHints(QtGui.QPainter.Antialiasing | QtGui.QPainter.SmoothPixmapTransform)
        self.setTransformationAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.setResizeAnchor(QtWidgets.QGraphicsView.AnchorUnderMouse)
        self.setDragMode(QtWidgets.QGraphicsView.ScrollHandDrag)
        self.scale_factor = 1.15

    def wheelEvent(self, event):
        zoom_in = event.angleDelta().y() > 0
        factor = self.scale_factor if zoom_in else 1 / self.scale_factor
        self.scale(factor, factor)

# Thiết kế Giao diện QT5
class Ui_MainWindow(object):

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1083, 845)
        MainWindow.setWindowTitle("PHẦN MỀM NHẬN DIỆN - ĐÁNH GIÁ VẾT NỨT/  HV: PHẠM NGỌC ĐĂNG KHOA")

        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setStyleSheet("font-size: 9pt;")  # tăng 30% từ mặc định 10pt
        self.centralwidget.setObjectName("centralwidget")

        self.base_width = 1083
        self.base_height = 845
        self.widgets = {}

        self.groupBox = QtWidgets.QGroupBox("Thông Tin Cấu Kiện:", self.centralwidget)
        self.groupBox.setGeometry(QtCore.QRect(0, 0, 350, 91))

        self.widgets['groupBox'] = self.groupBox

        self.label = QtWidgets.QLabel("Tên Cấu Kiện:", self.groupBox)
        self.label.setGeometry(QtCore.QRect(50, 30, 91, 16))
        self.label_2 = QtWidgets.QLabel("Ảnh gốc", self.groupBox)
        self.label_2.setGeometry(QtCore.QRect(130, 60, 55, 16))

        self.comboBox1 = QtWidgets.QComboBox(self.groupBox)
        self.comboBox1.setGeometry(QtCore.QRect(150, 28, 151, 22))
        self.comboBox1.addItems(["Dầm Bê Tông", "Cột Bê Tông", "Sàn Bê Tông", "Vách Bê Tông", "Tường Gạch"])

        self.chonanh1 = QtWidgets.QPushButton("Chọn Ảnh Tổng Thể", self.centralwidget)
        self.chonanh1.setGeometry(QtCore.QRect(800, 110, 150, 28))
        self.widgets['chonanh1'] = self.chonanh1

        self.anh1 = QtWidgets.QGraphicsView(self.centralwidget)
        self.anh1.setGeometry(QtCore.QRect(10, 80, 351, 351))
        self.widgets['anh1'] = self.anh1

        self.anh2 = QtWidgets.QGraphicsView(self.centralwidget)
        self.anh2.setGeometry(QtCore.QRect(370, 80, 351, 351))
        self.widgets['anh2'] = self.anh2

        self.anh3 = QtWidgets.QGraphicsView(self.centralwidget)
        self.anh3.setGeometry(QtCore.QRect(10, 440, 351, 351))
        self.widgets['anh3'] = self.anh3

        self.anh4 = QtWidgets.QGraphicsView(self.centralwidget)
        self.anh4.setGeometry(QtCore.QRect(370, 440, 351, 351))
        self.widgets['anh4'] = self.anh4

        self.label_3 = QtWidgets.QLabel("Kết Quả", self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(490, 60, 55, 16))
        self.widgets['label_3'] = self.label_3

        self.groupBox_2 = QtWidgets.QGroupBox("Đánh Giá:", self.centralwidget)
        self.groupBox_2.setGeometry(QtCore.QRect(730, 180, 341, 231))
        self.widgets['groupBox_2'] = self.groupBox_2

        self.lab2 = QtWidgets.QLabel("Mô tả Kết Quả", self.groupBox_2)
        self.lab2.setGeometry(QtCore.QRect(10, 30, 321, 191))
        self.lab2.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
        self.lab2.setWordWrap(True)

        self.groupBox_3 = QtWidgets.QGroupBox("Nhận xét:", self.centralwidget)
        self.groupBox_3.setGeometry(QtCore.QRect(730, 530, 341, 261))
        self.widgets['groupBox_3'] = self.groupBox_3

        self.lab3 = QtWidgets.QLabel("Mô tả Kết Quả", self.groupBox_3)
        self.lab3.setGeometry(QtCore.QRect(10, 20, 321, 231))
        self.lab3.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignLeft)
        self.lab3.setWordWrap(True)

        self.chonanh2 = QtWidgets.QPushButton("Chọn Ảnh Cận Cảnh", self.centralwidget)
        self.chonanh2.setGeometry(QtCore.QRect(800, 460, 150, 28))
        self.widgets['chonanh2'] = self.chonanh2

        # Thêm nhãn "Ghi chú:"
        self.label_input = QtWidgets.QLabel("Kích thước vât chuẩn:", self.centralwidget)
        self.label_input.setGeometry(QtCore.QRect(390, 30, 150, 20))

        self.label_input = QtWidgets.QLabel("x", self.centralwidget)
        self.label_input.setGeometry(QtCore.QRect(580, 28, 20, 20))

        self.label_input = QtWidgets.QLabel("mm", self.centralwidget)
        self.label_input.setGeometry(QtCore.QRect(630, 28, 20, 20))

        # Thêm ô nhập văn bản
        self.txtInput1 = QtWidgets.QLineEdit(self.centralwidget)
        self.txtInput1.setGeometry(QtCore.QRect(542, 28, 30, 22))  # x, y, width, height
        self.txtInput1.setPlaceholderText("20")

        self.txtInput2 = QtWidgets.QLineEdit(self.centralwidget)
        self.txtInput2.setGeometry(QtCore.QRect(600, 28, 30, 22))  # x, y, width, height
        self.txtInput2.setPlaceholderText("20")

        self.label_5 = QtWidgets.QLabel("Bước 1: Tìm Kiếm phát hiện vết nứt", self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(740, 80, 321, 16))
        self.widgets['label_5'] = self.label_5

        self.label_6 = QtWidgets.QLabel("Bước 2: Đánh Giá Vết nứt theo TCVN", self.centralwidget)
        self.label_6.setGeometry(QtCore.QRect(730, 440, 321, 16))
        self.widgets['label_6'] = self.label_6

        self.label_7 = QtWidgets.QLabel("Ghi chú: Chụp hình ảnh tổng thể cấu kiện.", self.centralwidget)
        self.label_7.setGeometry(QtCore.QRect(740, 150, 321, 16))
        self.widgets['label_7'] = self.label_7

        self.label_8 = QtWidgets.QLabel("Ghi chú: Chụp hình ảnh cận vết nứt", self.centralwidget)
        self.label_8.setGeometry(QtCore.QRect(740, 500, 321, 16))
        self.widgets['label_8'] = self.label_8



        # Kết xuất sang word
        self.export_button = QtWidgets.QPushButton("Kết xuất báo cáo Word", self.centralwidget)
        self.export_button.setGeometry(QtCore.QRect(800, 23, 180, 28))
        self.export_button.setObjectName("export_button")  # thêm dòng này

        # Nút luu tạm
        self.temp_save_button = QtWidgets.QPushButton("Lưu tạm kết quả", self.centralwidget)
        self.temp_save_button.setGeometry(QtCore.QRect(750, 750, 180, 28))
        self.widgets['temp_save_button'] = self.temp_save_button
        # Hiện số lần lưu
        self.lbl_save_count = QtWidgets.QLabel("Đã Lưu: 0", self.centralwidget)
        self.lbl_save_count.setGeometry(QtCore.QRect(950, 755, 180, 20))  # Tuỳ chỉnh vị trí nếu cần
        self.widgets['lbl_save_count'] = self.lbl_save_count
        # Xem biểu đồ
        self.show_chart_button = QtWidgets.QPushButton("Hiển thị Biểu đồ", self.centralwidget)
        self.show_chart_button.setGeometry(QtCore.QRect(800, 720, 180, 28))
        self.widgets['show_chart_button'] = self.show_chart_button



        self.original_positions = {
            name: widget.geometry()
            for name, widget in self.widgets.items()
        }

        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        MainWindow.setStatusBar(self.statusbar)

# Sự Kiện
class GlobalColorPicker(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(QtCore.Qt.FramelessWindowHint | QtCore.Qt.WindowStaysOnTopHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.setCursor(QtCore.Qt.CrossCursor)
        self.setGeometry(QtWidgets.QApplication.desktop().geometry())
        self.showFullScreen()


# Sự kiện
class MyMainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.base_width = self.ui.base_width
        self.base_height = self.ui.base_height

        self.ui.chonanh1.clicked.connect(self.handle_chonanh1)
        self.ui.anh1.mousePressEvent = self.open_zoomed_image1
        self.ui.anh2.mousePressEvent = self.open_zoomed_image2
        self.ui.chonanh2.clicked.connect(self.handle_chonanh2)
        self.ui.anh3.mousePressEvent = self.open_zoomed_image3
        self.ui.anh4.mousePressEvent = self.open_zoomed_image4
        # Sự kiện kết xuất word
        self.ui.export_button.clicked.connect(self.export_report_to_word)
        # Sự kiện lưu tạm
        self.cached_reports = []  # Lưu các lần chụp ảnh + kết quả
        self.ui.temp_save_button.clicked.connect(self.save_temporary_result)
        # Sự kiện đã lưu bao nhiêu lần
        self.temp_results = []
        self.update_save_count()
        #  Sự kiện nút hiện biểu đồ
        self.ui.show_chart_button.clicked.connect(self.show_crack_chart)

    def resizeEvent(self, event):
        w_ratio = self.width() / self.base_width
        h_ratio = self.height() / self.base_height

        for name, orig_geo in self.ui.original_positions.items():
            new_x = int(orig_geo.x() * w_ratio)
            new_y = int(orig_geo.y() * h_ratio)
            new_w = int(orig_geo.width() * w_ratio)
            new_h = int(orig_geo.height() * h_ratio)
            self.ui.widgets[name].setGeometry(new_x, new_y, new_w, new_h)

        super().resizeEvent(event)

    def handle_chonanh1(self):
            file_path, _ = QFileDialog.getOpenFileName(
                self, "Chọn ảnh để xử lý", "", "Image Files (*.jpg *.png *.jpeg *.bmp)"
            )
            if not file_path:
                return

            print("✅ Đã chọn ảnh:", file_path)

            pixmap = QtGui.QPixmap(file_path)
            self.image1 = cv2.imread(file_path)
            if pixmap.isNull():
                QtWidgets.QMessageBox.critical(self, "Lỗi", "Không thể đọc ảnh.")
                return

            # Hiển thị ảnh tổng thể lên ô anh1
            scaled_pixmap = pixmap.scaled(
                self.ui.anh1.width(), self.ui.anh1.height(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation
            )
            scene1 = QtWidgets.QGraphicsScene()
            scene1.addPixmap(scaled_pixmap)
            self.ui.anh1.setScene(scene1)

            # Hiện hộp thoại chờ
            self.progress_dialog = QtWidgets.QProgressDialog("⏳ Đang xử lý ảnh tổng thể...", None, 0, 0, self)
            self.progress_dialog.setWindowTitle("Thông báo")
            self.progress_dialog.setCancelButton(None)
            self.progress_dialog.setWindowModality(QtCore.Qt.ApplicationModal)
            self.progress_dialog.resize(300, 100)
            self.progress_dialog.show()
            QtWidgets.QApplication.processEvents()

            # Khởi tạo và chạy WorkerThread
            self.worker_thread1 = WorkerThread(
                image_path=file_path,
                selected_struct="",  # không dùng cho ảnh tổng thể
                real_width_cm=0,  # không dùng
                real_height_cm=0,  # không dùng
                mode="overview"  # để gọi run_pipeline
            )
            self.worker_thread1.result_ready.connect(self.on_overview_result_ready)
            self.worker_thread1.error_occurred.connect(self.on_crack_error)
            self.worker_thread1.start()

    def handle_chonanh2(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Chọn ảnh cận vết nứt", "", "Image Files (*.jpg *.png *.jpeg *.bmp)"
        )
        if not file_path:
            return

        pixmap = QtGui.QPixmap(file_path)
        if pixmap.isNull():
            QtWidgets.QMessageBox.critical(self, "Lỗi", "Không thể đọc ảnh.")
            return

        # Hiển thị ảnh gốc vào anh3
        scaled_pixmap = pixmap.scaled(
            self.ui.anh3.width(), self.ui.anh3.height(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation
        )
        scene3 = QtWidgets.QGraphicsScene()
        scene3.addPixmap(scaled_pixmap)
        self.ui.anh3.setScene(scene3)
        self.image3 = cv2.imread(file_path)

        # Lấy nội dung vật chuẩn
        selected_struct = self.ui.comboBox1.currentText()
        try:
            text_w = self.ui.txtInput1.text().strip()
            text_h = self.ui.txtInput2.text().strip()
            real_width_cm = float(text_w) if text_w else 20
            real_height_cm = float(text_h) if text_h else 20
        except ValueError:
            QtWidgets.QMessageBox.warning(self, "Lỗi", "Vui lòng nhập kích thước hợp lệ")
            return

        # Hiện hộp thoại tiến trình
        self.progress_dialog = QtWidgets.QProgressDialog("⏳ Đang xử lý ảnh cận...", None, 0, 0, self)
        self.progress_dialog.setWindowTitle("Thông báo")
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setWindowModality(QtCore.Qt.ApplicationModal)
        self.progress_dialog.resize(300, 100)
        self.progress_dialog.show()
        QtWidgets.QApplication.processEvents()

        # Khởi động worker thread
        self.worker_thread2 = WorkerThread(
            image_path=file_path,
            selected_struct=selected_struct,
            real_width_cm=real_width_cm,
            real_height_cm=real_height_cm,
            mode="zoom"  # gọi run_pipeline2
        )
        self.worker_thread2.result_ready.connect(self.on_crack_result_ready)
        self.worker_thread2.error_occurred.connect(self.on_crack_error)
        self.worker_thread2.start()

    # Xử lý đa luồn ghi kết quả
    def on_crack_result_ready(self, result_image, result_text, chart_np):

        self.chart_np = chart_np

        self.progress_dialog.close()

        self.image4 = result_image
        height, width, channel = result_image.shape
        bytes_per_line = 3 * width
        qimage = QtGui.QImage(result_image.data, width, height, bytes_per_line, QtGui.QImage.Format_BGR888)
        pixmap_result = QtGui.QPixmap.fromImage(qimage)

        scaled_result = pixmap_result.scaled(
            self.ui.anh4.width(), self.ui.anh4.height(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation
        )
        scene4 = QtWidgets.QGraphicsScene()
        scene4.addPixmap(scaled_result)
        self.ui.anh4.setScene(scene4)

        self.ui.lab3.setText(result_text)

    # Xử lý đa luồn ghi kết quả
    def on_overview_result_ready(self, result_image, result_text):
        self.progress_dialog.close()

        self.image2 = result_image
        height, width, channel = result_image.shape
        bytes_per_line = 3 * width
        qimage = QtGui.QImage(result_image.data, width, height, bytes_per_line, QtGui.QImage.Format_BGR888)
        pixmap_result = QtGui.QPixmap.fromImage(qimage)

        scaled_result = pixmap_result.scaled(
            self.ui.anh2.width(), self.ui.anh2.height(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation
        )
        scene2 = QtWidgets.QGraphicsScene()
        scene2.addPixmap(scaled_result)
        self.ui.anh2.setScene(scene2)

        self.ui.lab2.setText(result_text)

    def on_crack_error(self, error_msg):
        self.progress_dialog.close()
        QtWidgets.QMessageBox.critical(self, "Lỗi xử lý ảnh cận:", error_msg)

    def open_zoomed_image(self, image, title="Ảnh"):
        if image is None:
            QtWidgets.QMessageBox.information(self, "Thông báo", "Chưa có ảnh để hiển thị.")
            return

        height, width, channel = image.shape
        bytes_per_line = 3 * width
        qimage = QtGui.QImage(image.data, width, height, bytes_per_line, QtGui.QImage.Format_BGR888)
        pixmap = QtGui.QPixmap.fromImage(qimage)

        viewer = ZoomableImageViewer(pixmap)

        dialog = QtWidgets.QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setFixedSize(640, 640)

        layout = QtWidgets.QVBoxLayout(dialog)
        layout.addWidget(viewer)
        dialog.setLayout(layout)
        dialog.exec_()

    def open_zoomed_image1(self, event):
        self.open_zoomed_image(getattr(self, 'image1', None), "Ảnh gốc")

    def open_zoomed_image2(self, event):
        self.open_zoomed_image(getattr(self, 'image2', None), "Ảnh kết quả")

    def open_zoomed_image3(self, event):
        self.open_zoomed_image(getattr(self, 'image3', None), "Ảnh cận gốc")

    def open_zoomed_image4(self, event):
        self.open_zoomed_image(getattr(self, 'image4', None), "Kết quả phân tích")

    def export_report_to_word(self):
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from bs4 import BeautifulSoup
            import tempfile
            import os
            import cv2

            def add_html_text(paragraph, html_text):
                soup = BeautifulSoup(html_text.replace("<br>", "\n"), "html.parser")
                for item in soup.contents:
                    if item.name == "span":
                        run = paragraph.add_run(item.get_text())
                        style = item.get("style", "").lower()
                        if "bold" in style:
                            run.bold = True
                        if "color:red" in style:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    elif isinstance(item, str):
                        paragraph.add_run(item)
                    elif item.name == "br":
                        paragraph.add_run().add_break()
                    else:
                        paragraph.add_run(item.get_text())

            if not self.cached_reports:
                QtWidgets.QMessageBox.warning(self, "Cảnh báo", "Chưa có nội dung nào được lưu tạm.")
                return

            doc = Document()
            doc.add_heading("BÁO CÁO NHẬN DIỆN VÀ ĐÁNH GIÁ", 0)

            for idx, report in enumerate(self.cached_reports, 1):
                doc.add_heading(f"I.{idx} Kết quả lần {idx}", level=1)
                doc.add_paragraph("1. Thông tin cấu kiện", style='Heading 2')
                doc.add_paragraph(f"Tên cấu kiện: {report['ten_cau_kien']}")
                doc.add_paragraph("2. Hình ảnh cấu kiện", style='Heading 2')

                table = doc.add_table(rows=4, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                images = [
                    ('image1', "Hình 1: Ảnh tổng thể cấu kiện trước xử lý"),
                    ('image2', "Hình 2: Ảnh phát hiện các vết nứt"),
                    ('image3', "Hình 3: Ảnh cận cấu kiện trước xử lý"),
                    ('image4', "Hình 4: Ảnh phát hiện vết nứt cận cảnh"),
                ]

                row = 0
                col = 0
                for img_attr, caption in images:
                    img = report.get(img_attr)
                    if img is not None:
                        try:
                            temp_path = os.path.join(tempfile.gettempdir(), f"{img_attr}_{idx}.png")
                            cv2.imwrite(temp_path, img)
                            img_cell = table.cell(row, col)
                            img_paragraph = img_cell.paragraphs[0]
                            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            img_paragraph.add_run().add_picture(temp_path, width=Inches(2.8))

                            caption_cell = table.cell(row + 1, col)
                            caption_para = caption_cell.paragraphs[0]
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.add_run(caption).italic = True
                        except Exception as e:
                            print(f"Lỗi lưu ảnh {img_attr}: {e}")
                    col += 1
                    if col >= 2:
                        col = 0
                        row += 2
                chart_np = report.get("chart_np")
                if chart_np is not None:
                    try:
                        import tempfile
                        chart_path = os.path.join(tempfile.gettempdir(), f"chart_{idx}.png")
                        cv2.imwrite(chart_path, chart_np)

                        doc.add_paragraph("3. Biểu đồ chiều rộng vết nứt:", style='Heading 2')
                        doc.add_picture(chart_path, width=Inches(5.5))
                    except Exception as e:
                        print(f"Lỗi lưu biểu đồ: {e}")
                doc.add_paragraph("4. Số lượng vết nứt tổng thể:", style='Heading 2')
                doc.add_paragraph(report.get('lab2_text', ''))

                doc.add_paragraph("5. Kết quả vết nứt ảnh cận:", style='Heading 2')

                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Inches(0.3)
                add_html_text(para, report.get('lab3_text', ''))


            save_path, _ = QFileDialog.getSaveFileName(self, "Lưu báo cáo", "", "Word Documents (*.docx)")
            if save_path:
                if not save_path.endswith(".docx"):
                    save_path += ".docx"
                doc.save(save_path)
                QtWidgets.QMessageBox.information(self, "Thành công", f"Đã lưu báo cáo tại:\n{save_path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Xuất báo cáo thất bại:\n{str(e)}")

    # hàm lưu tạm
    def save_temporary_result(self):
        report = {
            "ten_cau_kien": self.ui.comboBox1.currentText(),
            "lab2_text": self.ui.lab2.text(),
            "lab3_text": self.ui.lab3.text(),
            "image1": getattr(self, "image1", None),
            "image2": getattr(self, "image2", None),
            "image3": getattr(self, "image3", None),
            "image4": getattr(self, "image4", None),
            "chart_np": getattr(self, "chart_np", None),  # thêm dòng này
        }
        self.cached_reports.append(report)
        self.update_save_count()

    # Cập nhật số lần lưu tạm
    def update_save_count(self):
        self.ui.lbl_save_count.setText(f"Đã Lưu: {len(self.cached_reports)}")
    # Hiện biểu đồ
    def show_crack_chart(self):
        if not hasattr(self, "chart_np") or self.chart_np is None:
            QtWidgets.QMessageBox.warning(self, "Thông báo", "Chưa có biểu đồ nào được tạo.")
            return

        try:
            # Tạo QImage từ NumPy array
            img = self.chart_np.copy()
            height, width, channel = img.shape
            bytes_per_line = 3 * width
            qimage = QtGui.QImage(img.data, width, height, bytes_per_line, QtGui.QImage.Format_BGR888)
            pixmap = QtGui.QPixmap.fromImage(qimage)

            # Tạo viewer có thể zoom
            viewer = ZoomableImageViewer(pixmap)

            dialog = QtWidgets.QDialog(self)
            dialog.setWindowTitle("Biểu đồ chiều rộng vết nứt")
            dialog.setFixedSize(1200, 800)

            layout = QtWidgets.QVBoxLayout(dialog)
            layout.addWidget(viewer)
            dialog.setLayout(layout)
            dialog.exec_()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Lỗi", f"Lỗi khi hiển thị biểu đồ:\n{str(e)}")


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    window = MyMainWindow()
    window.show()
    sys.exit(app.exec_())